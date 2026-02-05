from pathlib import Path
from typing import Optional
import re
import csv
import json

# Check for OCR support
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

# Check for PDF to image conversion
try:
    from pdf2image import convert_from_path
    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False

class Parser:
    def __init__(self):
        self.current_file: Optional[Path] = None
        self.content: str = ""
        self.current_title: str = ""  # Store extracted title
    
    def extract_title(self, file_path: Path) -> str:
        """Extract the title from a document file"""
        self.current_file = file_path
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.txt':
                return self._extract_title_txt(file_path)
            elif extension in ['.docx', '.doc']:
                return self._extract_title_docx(file_path)
            elif extension == '.pdf':
                return self._extract_title_pdf(file_path)
            elif extension in ['.md', '.markdown']:
                return self._extract_title_markdown(file_path)
            elif extension in ['.html', '.htm']:
                return self._extract_title_html(file_path)
            elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif']:
                return self._extract_title_image(file_path)
            else:
                # For unsupported formats, return filename
                return file_path.stem
        except Exception as e:
            print(f"Error extracting title from {file_path}: {e}")
            return file_path.stem
    
    def parse_file(self, file_path: Path) -> str:
        
        self.current_file = file_path
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.txt':
                return self._parse_txt(file_path)
            elif extension in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            elif extension == '.pdf':
                return self._parse_pdf(file_path)
            elif extension in ['.xlsx', '.xls']:
                return self._parse_excel(file_path)
            elif extension in ['.csv', '.tsv']:
                return self._parse_csv(file_path, extension)
            elif extension == '.rtf':
                return self._parse_rtf(file_path)
            elif extension == '.odt':
                return self._parse_odt(file_path)
            elif extension in ['.md', '.markdown']:
                return self._parse_txt(file_path)  # Markdown is plain text
            elif extension in ['.html', '.htm']:
                return self._parse_html(file_path)
            elif extension == '.xml':
                return self._parse_xml(file_path)
            elif extension == '.json':
                return self._parse_json(file_path)
            elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif']:
                return self._parse_image(file_path)
            else:
                return ""
        except Exception as e:
            print(f"Error parsing {file_path}: {e}")
            return ""
    
    def _parse_txt(self, file_path: Path) -> str:
        
        # Try common encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"Error reading {file_path} with {encoding}: {e}")
                continue
        
        return ""
    
    def _parse_docx(self, file_path: Path) -> str:
        
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            full_text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return ""
        except Exception as e:
            print(f"Error parsing DOCX {file_path}: {e}")
            return ""
    
    def _parse_pdf(self, file_path: Path) -> str:
        
        try:
            from PyPDF2 import PdfReader
            
            # Validate file exists and is not empty
            if not file_path.exists():
                print(f"PDF file does not exist: {file_path}")
                return ""
            
            if file_path.stat().st_size == 0:
                print(f"PDF file is empty: {file_path}")
                return ""
            
            reader = PdfReader(str(file_path))
            full_text = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text)
            
            extracted_text = '\n'.join(full_text)
            
            # If text extraction yielded little/no text, try OCR
            if len(extracted_text.strip()) < 100 and HAS_OCR and HAS_PDF2IMAGE:
                print(f"PDF appears to be scanned, attempting OCR: {file_path.name}")
                ocr_text = self._ocr_pdf(file_path)
                if ocr_text:
                    return ocr_text
            
            return extracted_text
        except ImportError:
            print("PyPDF2 not installed. Install with: pip install PyPDF2")
            return ""
        except Exception as e:
            print(f"Error parsing PDF {file_path}: {e}")
            return ""
    
    def _parse_excel(self, file_path: Path) -> str:
        try:
            import openpyxl
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            full_text = []
            
            for sheet in workbook.worksheets:
                full_text.append(f"Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' '.join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        full_text.append(row_text)
            
            return '\n'.join(full_text)
        except ImportError:
            print("openpyxl not installed. Install with: pip install openpyxl")
            return ""
        except Exception as e:
            print(f"Error parsing Excel {file_path}: {e}")
            return ""
    
    def _parse_csv(self, file_path: Path, extension: str) -> str:
        try:
            delimiter = '\t' if extension == '.tsv' else ','
            full_text = []
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f, delimiter=delimiter)
                for row in reader:
                    row_text = ' '.join(str(cell) for cell in row if cell)
                    if row_text.strip():
                        full_text.append(row_text)
            
            return '\n'.join(full_text)
        except Exception as e:
            print(f"Error parsing CSV/TSV {file_path}: {e}")
            return ""
    
    def _parse_rtf(self, file_path: Path) -> str:
        try:
            # Basic RTF parsing - strip RTF control words
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Remove RTF control words (basic approach)
            content = re.sub(r'\\[a-z]+\d*\s?', ' ', content)
            content = re.sub(r'[{}]', '', content)
            content = re.sub(r'\s+', ' ', content)
            
            return content.strip()
        except Exception as e:
            print(f"Error parsing RTF {file_path}: {e}")
            return ""
    
    def _parse_odt(self, file_path: Path) -> str:
        try:
            from zipfile import ZipFile
            from xml.etree import ElementTree as ET
            
            full_text = []
            with ZipFile(file_path) as odt_file:
                content_xml = odt_file.read('content.xml')
                tree = ET.fromstring(content_xml)
                
                # Extract all text nodes
                for elem in tree.iter():
                    if elem.text:
                        full_text.append(elem.text)
                    if elem.tail:
                        full_text.append(elem.tail)
            
            return '\\n'.join(full_text)
        except Exception as e:
            print(f"Error parsing ODT {file_path}: {e}")
            return ""
    
    def _parse_html(self, file_path: Path) -> str:
        try:
            from html.parser import HTMLParser
            
            class TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text = []
                
                def handle_data(self, data):
                    self.text.append(data)
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            parser = TextExtractor()
            parser.feed(content)
            
            return ' '.join(parser.text)
        except Exception as e:
            print(f"Error parsing HTML {file_path}: {e}")
            return ""
    
    def _parse_xml(self, file_path: Path) -> str:
        try:
            from xml.etree import ElementTree as ET
            
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            full_text = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    full_text.append(elem.text)
                if elem.tail and elem.tail.strip():
                    full_text.append(elem.tail)
            
            return '\\n'.join(full_text)
        except Exception as e:
            print(f"Error parsing XML {file_path}: {e}")
            return ""
    
    def _parse_json(self, file_path: Path) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Extract all string values recursively
            def extract_strings(obj, strings=[]):
                if isinstance(obj, dict):
                    for value in obj.values():
                        extract_strings(value, strings)
                elif isinstance(obj, list):
                    for item in obj:
                        extract_strings(item, strings)
                elif isinstance(obj, str):
                    strings.append(obj)
                return strings
            
            text_data = extract_strings(data)
            return '\\n'.join(text_data)
        except Exception as e:
            print(f"Error parsing JSON {file_path}: {e}")
            return ""
    
    @staticmethod
    def clean_text(text: str) -> str:
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters
        text = ''.join(char for char in text if ord(char) >= 32 or char == '\n')
        
        return text.strip()
    
    @staticmethod
    def split_into_sentences(text: str) -> list:
        
        # Simple sentence splitting
        sentences = re.split(r'[.!?]+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _extract_title_txt(self, file_path: Path) -> str:
        """Extract title from text file (first non-empty line)"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            # Remove markdown heading markers if present
                            if line.startswith('#'):
                                line = line.lstrip('#').strip()
                            # Return first non-empty line, truncate if too long
                            return line[:100] if len(line) > 100 else line
                return file_path.stem
            except:
                continue
        return file_path.stem
    
    def _extract_title_markdown(self, file_path: Path) -> str:
        """Extract title from markdown file (first heading or first line)"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            # Check for markdown heading
                            if line.startswith('#'):
                                title = line.lstrip('#').strip()
                                return title[:100] if len(title) > 100 else title
                            # Otherwise return first line
                            return line[:100] if len(line) > 100 else line
                return file_path.stem
            except:
                continue
        return file_path.stem
    
    def _extract_title_docx(self, file_path: Path) -> str:
        """Extract title from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            
            # Try to get title from core properties first
            if doc.core_properties.title:
                return doc.core_properties.title
            
            # Otherwise get first paragraph
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    return text[:100] if len(text) > 100 else text
            
            return file_path.stem
        except ImportError:
            return file_path.stem
        except Exception as e:
            print(f"Error extracting DOCX title: {e}")
            return file_path.stem
    
    def _extract_title_pdf(self, file_path: Path) -> str:
        """Extract title from PDF file"""
        try:
            from PyPDF2 import PdfReader
            
            # Validate file exists and is not empty
            if not file_path.exists():
                print(f"PDF file does not exist: {file_path}")
                return file_path.stem
            
            if file_path.stat().st_size == 0:
                print(f"PDF file is empty: {file_path}")
                return file_path.stem
            
            reader = PdfReader(str(file_path))
            
            # Try to get title from metadata first (but skip generic/useless titles)
            if reader.metadata and reader.metadata.title:
                metadata_title = reader.metadata.title.strip()
                # Skip generic placeholder titles
                generic_titles = [
                    'untitled', 'untitled document', 'document',  
                    'document1', 'new document', 'doc1', 'doc',
                    file_path.stem,  # Skip if it's just the filename
                ]
                if metadata_title.lower() not in generic_titles and len(metadata_title) > 3:
                    return metadata_title
            
            # Otherwise get first meaningful line from first page
            if reader.pages:
                text = reader.pages[0].extract_text()
                if text:
                    lines = text.split('\n')
                    
                    # Skip common non-title patterns
                    skip_patterns = [
                        r'^page\s+\d+',  # "Page 1", "Page 1 of 19"
                        r'^\d+\s+of\s+\d+',  # "1 of 19"
                        r'^©',  # Copyright
                        r'^all rights reserved',
                        r'^confidential',
                        r'^draft',
                        r'^\d{1,4}[-/]\d{1,2}[-/]\d{1,4}',  # Dates
                        r'^[A-Z\s]{20,}$',  # ALL CAPS headers (likely column headers) - only uppercase
                        r'^(name|code|branch|address|phone|email|date|description|title|id|number)[\s,|]+',  # Table column headers
                        r'\b\d{2,}\s+\d{2,}',  # Multiple numbers separated by spaces (likely data rows)
                        r'\b(ltd|limited|inc|corp|corporation)\s+\d',  # Company name followed by numbers (likely data)
                    ]
                    
                    for line in lines:
                        line = line.strip()
                        if not line or len(line) < 3:
                            continue
                        
                        # Check if line is all caps (column headers are often all caps)
                        if len(line) >= 20 and line.replace(' ', '').isupper():
                            continue
                        
                        # Skip lines with multiple numbers (likely data rows, not titles)
                        numbers_in_line = len(re.findall(r'\d+', line))
                        if numbers_in_line >= 3:
                            continue
                        
                        # Check if line matches skip patterns
                        should_skip = False
                        for pattern in skip_patterns:
                            if re.search(pattern, line, re.IGNORECASE):
                                should_skip = True
                                break
                        
                        if not should_skip:
                            # Found a good title candidate
                            return line[:100] if len(line) > 100 else line
            
            return file_path.stem
        except ImportError:
            return file_path.stem
        except Exception as e:
            print(f"Error extracting PDF title from {file_path.name}: {e}")
            return file_path.stem
    
    def _extract_title_html(self, file_path: Path) -> str:
        """Extract title from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Try to find <title> tag
            title_match = re.search(r'<title[^>]*>(.*?)</title>', content, re.IGNORECASE | re.DOTALL)
            if title_match:
                title = title_match.group(1).strip()
                # Clean HTML entities and tags
                title = re.sub(r'<[^>]+>', '', title)
                title = re.sub(r'&[a-zA-Z]+;', ' ', title)
                return title[:100] if len(title) > 100 else title
            
            # Try to find first h1 tag
            h1_match = re.search(r'<h1[^>]*>(.*?)</h1>', content, re.IGNORECASE | re.DOTALL)
            if h1_match:
                title = h1_match.group(1).strip()
                title = re.sub(r'<[^>]+>', '', title)
                return title[:100] if len(title) > 100 else title
            
            return file_path.stem
        except Exception as e:
            print(f"Error extracting HTML title: {e}")
            return file_path.stem
    
    def _extract_title_image(self, file_path: Path) -> str:
        """Extract title from image file using OCR"""
        if not HAS_OCR:
            return file_path.stem
        
        try:
            from PIL import Image
            
            # OCR the image
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            
            if text and text.strip():
                # Get first non-empty line as title
                for line in text.split('\n'):
                    line = line.strip()
                    if line:
                        return line[:100] if len(line) > 100 else line
            
            return file_path.stem
        except Exception as e:
            print(f"Error extracting title from image {file_path.name}: {e}")
            return file_path.stem
    
    def _parse_image(self, file_path: Path) -> str:
        """Parse text from image file using OCR"""
        if not HAS_OCR:
            print(f"Warning: pytesseract not installed. Cannot OCR {file_path.name}")
            return ""
        
        try:
            from PIL import Image
            
            # Open and OCR the image
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            
            if text and len(text.strip()) > 0:
                print(f"OCR extracted {len(text)} characters from {file_path.name}")
            else:
                print(f"Warning: OCR returned no text from {file_path.name}")
            
            return text
        except Exception as e:
            print(f"Error during OCR of {file_path.name}: {e}")
            return ""
    
    def _ocr_pdf(self, file_path: Path) -> str:
        """Convert PDF pages to images and OCR them"""
        if not HAS_OCR or not HAS_PDF2IMAGE:
            print(f"Warning: pytesseract or pdf2image not installed. Cannot OCR {file_path.name}")
            return ""
        
        try:
            from pdf2image import convert_from_path
            from PIL import Image
            
            # Convert PDF pages to images
            print(f"Converting {file_path.name} to images for OCR...")
            images = convert_from_path(file_path)
            
            # OCR each page
            all_text = []
            for i, img in enumerate(images, 1):
                print(f"OCR'ing page {i}/{len(images)}...")
                page_text = pytesseract.image_to_string(img)
                if page_text.strip():
                    all_text.append(page_text)
            
            combined_text = "\n\n".join(all_text)
            print(f"OCR extracted {len(combined_text)} characters from {len(images)} pages")
            
            return combined_text
        except Exception as e:
            print(f"Error during PDF OCR of {file_path.name}: {e}")
            return ""
